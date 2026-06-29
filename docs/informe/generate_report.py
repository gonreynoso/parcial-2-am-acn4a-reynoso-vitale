#!/usr/bin/env python3
"""Genera el informe Word de Sprint (RunTracker) con capturas de pantalla."""

import os
import subprocess
import sys
from pathlib import Path

BASE = Path(__file__).parent
MOCKUPS = BASE / "mockups"
SCREENSHOTS = BASE / "screenshots"
OUTPUT = BASE / "Informe_Sprint_RunTracker.docx"

# ── Estilos compartidos ──────────────────────────────────────────────────────

COMMON_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
* { margin: 0; padding: 0; box-sizing: border-box; }
body {
  font-family: 'Inter', -apple-system, sans-serif;
  background: #0a0a0a;
  display: flex; justify-content: center; align-items: flex-start;
  min-height: 100vh; padding: 20px;
}
.phone {
  width: 390px; height: 844px;
  background: #121212;
  border-radius: 40px;
  overflow: hidden;
  position: relative;
  box-shadow: 0 0 0 3px #2a2a2a, 0 25px 80px rgba(0,230,168,0.08);
}
.status-bar {
  height: 44px; display: flex; align-items: center; justify-content: space-between;
  padding: 0 28px; color: #fff; font-size: 14px; font-weight: 600;
}
.content { padding: 0 20px 100px; overflow-y: auto; height: calc(100% - 44px); }
.card {
  background: #1E1E1E; border-radius: 16px; padding: 16px; margin-bottom: 12px;
}
.neon { color: #00E6A8; }
.muted { color: #AAAAAA; }
.white { color: #fff; }
.bold { font-weight: 700; }
.navbar {
  position: absolute; bottom: 0; left: 0; right: 0;
  height: 72px; background: #1A1D21;
  display: flex; align-items: center; justify-content: space-around;
  border-top: 1px solid #2a2a2a; padding-bottom: 8px;
}
.nav-item { display: flex; flex-direction: column; align-items: center; gap: 4px; }
.nav-item span { font-size: 11px; }
.nav-item.active span, .nav-item.active svg { color: #00E6A8; }
.nav-item.inactive span, .nav-item.inactive svg { color: #8A8F98; }
.fab {
  width: 56px; height: 56px; background: #00E6A8; border-radius: 16px;
  display: flex; align-items: center; justify-content: center;
  margin-top: -28px; box-shadow: 0 4px 20px rgba(0,230,168,0.4);
}
.metric-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 8px; }
.metric-card {
  background: #1E1E1E; border-radius: 16px; padding: 16px;
  text-align: center;
}
.metric-value { font-size: 20px; font-weight: 700; color: #fff; margin: 8px 0 4px; }
.metric-label { font-size: 12px; color: #888; }
.progress-bar {
  height: 8px; background: #333; border-radius: 4px; margin-top: 10px; overflow: hidden;
}
.progress-fill { height: 100%; background: #00E6A8; border-radius: 4px; }
.btn-primary {
  background: #00E6A8; color: #000; border: none; border-radius: 8px;
  padding: 14px; width: 100%; font-size: 15px; font-weight: 600; text-align: center;
}
.input {
  background: #111; border: 1px solid #3A3A3A; border-radius: 8px;
  padding: 14px; color: #fff; width: 100%; font-size: 14px; margin-top: 6px;
}
.btn-outline {
  border: 1px solid #3A3A3A; border-radius: 8px; padding: 12px;
  color: #fff; text-align: center; flex: 1; font-size: 13px;
}
"""

SCREENS = {
    "01_splash": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.splash-content {
  display: flex; flex-direction: column; align-items: center; justify-content: center;
  height: 100%; text-align: center;
}
.logo {
  width: 80px; height: 80px; background: #00E6A8; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-size: 36px; font-weight: 800; color: #000;
  box-shadow: 0 0 40px rgba(0,230,168,0.3);
}
.app-name { font-size: 32px; font-weight: 800; color: #fff; margin-top: 20px; }
.tagline { color: #888; margin-top: 8px; font-size: 14px; }
</style></head><body>
<div class="phone"><div class="splash-content">
  <div class="logo">S</div>
  <div class="app-name">Sprint</div>
  <div class="tagline">Tu compañero de carrera</div>
</div></div></body></html>""",

    "02_login": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.content { padding: 40px 32px; }
.brand { display: flex; align-items: center; gap: 10px; margin-bottom: 36px; }
.brand-logo {
  width: 32px; height: 32px; background: #00E6A8; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-weight: 800; color: #000; font-size: 13px;
}
.title { font-size: 22px; font-weight: 700; color: #fff; margin-bottom: 6px; }
.toggle { color: #888; font-size: 14px; margin-bottom: 28px; }
.social { display: flex; gap: 8px; margin-bottom: 20px; }
.or-row { display: flex; align-items: center; gap: 14px; margin-bottom: 20px; }
.or-line { flex: 1; height: 1px; background: #2A2A2A; }
.or-text { color: #888; font-size: 11px; letter-spacing: 1px; }
.label { color: #fff; font-size: 14px; margin-top: 16px; }
.forgot { color: #888; font-size: 13px; margin-top: 16px; }
</style></head><body>
<div class="phone">
  <div class="content">
    <div class="brand"><div class="brand-logo">S</div><span class="white bold" style="font-size:18px">Sprint</span></div>
    <div class="title">Bienvenido de vuelta</div>
    <div class="toggle">¿No tenés cuenta? Registrate</div>
    <div class="social">
      <div class="btn-outline">GitHub</div>
      <div class="btn-outline">Google</div>
    </div>
    <div class="or-row"><div class="or-line"></div><div class="or-text">OR</div><div class="or-line"></div></div>
    <div class="label">Email</div>
    <div class="input">gonzalo@email.com</div>
    <div class="label">Contraseña</div>
    <div class="input">••••••••</div>
    <div style="margin-top:16px" class="btn-primary">Ingresar</div>
    <div class="forgot">¿Olvidaste tu contraseña?</div>
  </div>
</div></body></html>""",

    "03_dashboard": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.steps-circle {
  width: 180px; height: 180px; border: 3px solid #00E6A8; border-radius: 50%;
  display: flex; flex-direction: column; align-items: center; justify-content: center;
  margin: 24px auto;
}
.steps-num { font-size: 48px; font-weight: 800; color: #00E6A8; }
.hoy { text-align: center; color: #fff; font-size: 18px; font-weight: 600; margin-top: 16px; }
.greeting { font-size: 22px; font-weight: 700; color: #fff; }
.frase { color: #888; font-size: 14px; margin-top: 4px; margin-bottom: 8px; }
.streak-row { display: flex; align-items: center; gap: 12px; }
.streak-emoji { font-size: 32px; }
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content" style="padding-bottom:90px">
    <div class="greeting">Hola, Gonzalo 👋</div>
    <div class="frase">Hoy es un buen día para correr</div>
    <div class="hoy">Hoy</div>
    <div class="steps-circle">
      <div class="steps-num">4,267</div>
      <div class="muted">Pasos</div>
    </div>
    <div class="card">
      <div style="display:flex;justify-content:space-between">
        <span class="white bold">Objetivo diario</span>
        <span class="neon bold">4,267 / 10,000</span>
      </div>
      <div class="progress-bar"><div class="progress-fill" style="width:43%"></div></div>
    </div>
    <div class="card streak-row">
      <div class="streak-emoji">🔥</div>
      <div>
        <div class="white bold">Racha de 5 días</div>
        <div class="muted" style="font-size:13px">¡Seguí así para mantenerla!</div>
      </div>
    </div>
    <div class="card">
      <div style="display:flex;align-items:center;gap:10px;margin-bottom:10px">
        <span style="font-size:20px">🌤️</span>
        <span class="white bold">Pronóstico y Actividad</span>
      </div>
      <div class="muted" style="font-size:14px">¿Salimos a correr o nos quedamos en casa?</div>
      <div class="neon bold" style="text-align:right;margin-top:10px;font-size:14px">VER CLIMA ➔</div>
    </div>
    <div class="metric-grid">
      <div class="metric-card"><div>📍</div><div class="metric-value">3.20 km</div><div class="metric-label">Distancia</div></div>
      <div class="metric-card"><div>🔥</div><div class="metric-value">192 cal</div><div class="metric-label">Calorías</div></div>
      <div class="metric-card"><div>⏱️</div><div class="metric-value">28:45</div><div class="metric-label">Tiempo</div></div>
      <div class="metric-card"><div>🏃</div><div class="metric-value">2</div><div class="metric-label">Carreras</div></div>
    </div>
  </div>
  <div class="navbar">
    <div class="nav-item active"><span>🏠</span><span>Inicio</span></div>
    <div class="fab">🏃</div>
    <div class="nav-item inactive"><span>👤</span><span>Perfil</span></div>
  </div>
</div></body></html>""",

    "04_perfil": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.avatar {
  width: 64px; height: 64px; background: #00E6A8; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
  font-size: 28px; font-weight: 800; color: #000;
}
.header-card { display: flex; align-items: center; gap: 16px; }
.stats-row { display: flex; justify-content: space-around; text-align: center; margin: 16px 0; }
.stat-val { font-size: 20px; font-weight: 700; color: #fff; }
.stat-lbl { font-size: 12px; color: #888; }
.achievement { display: flex; gap: 12px; }
.ach-item {
  flex: 1; background: #1E1E1E; border-radius: 12px; padding: 12px; text-align: center;
}
.ach-item.locked { opacity: 0.35; }
.run-card {
  background: #1E1E1E; border-radius: 12px; padding: 14px;
  display: flex; align-items: center; gap: 12px; margin-bottom: 8px;
}
.run-icon {
  width: 40px; height: 40px; background: #00E6A820; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
}
.section-title { font-size: 16px; font-weight: 700; color: #fff; margin: 16px 0 8px; }
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content" style="padding-bottom:90px">
    <div class="card header-card">
      <div class="avatar">G</div>
      <div>
        <div class="white bold" style="font-size:20px">Gonzalo</div>
        <div class="muted" style="font-size:13px">Te uniste en 2025</div>
      </div>
      <div style="margin-left:auto;display:flex;gap:12px">
        <span>✏️</span><span>🚪</span>
      </div>
    </div>
    <div class="card">
      <div class="muted" style="font-size:13px;margin-bottom:8px">Peso · Altura · IMC</div>
      <div style="display:flex;justify-content:space-between">
        <span class="white">72 kg</span><span class="white">178 cm</span><span class="neon bold">22.7</span>
      </div>
    </div>
    <div class="section-title">Estadísticas</div>
    <div class="card stats-row">
      <div><div class="stat-val">42.5 km</div><div class="stat-lbl">Km totales</div></div>
      <div><div class="stat-val">18</div><div class="stat-lbl">Carreras</div></div>
      <div><div class="stat-val">5:12</div><div class="stat-lbl">Mejor ritmo</div></div>
    </div>
    <div class="section-title">Logros</div>
    <div class="achievement">
      <div class="ach-item"><div>🏅</div><div class="white" style="font-size:12px;margin-top:4px">Primera carrera</div></div>
      <div class="ach-item"><div>🎯</div><div class="white" style="font-size:12px;margin-top:4px">5K</div></div>
      <div class="ach-item locked"><div>🏆</div><div class="white" style="font-size:12px;margin-top:4px">10K</div></div>
    </div>
    <div class="section-title">Mis carreras</div>
    <div class="run-card">
      <div class="run-icon">🏃</div>
      <div style="flex:1">
        <div class="white bold">28/06/2026</div>
        <div class="muted" style="font-size:12px">5:45 /km · 32:10</div>
      </div>
      <div class="neon bold">5.60 km</div>
    </div>
    <div class="run-card">
      <div class="run-icon">🚶</div>
      <div style="flex:1">
        <div class="white bold">27/06/2026</div>
        <div class="muted" style="font-size:12px">8:20 /km · 45:00</div>
      </div>
      <div class="neon bold">5.40 km</div>
    </div>
  </div>
  <div class="navbar">
    <div class="nav-item inactive"><span>🏠</span><span>Inicio</span></div>
    <div class="fab">🏃</div>
    <div class="nav-item active"><span>👤</span><span>Perfil</span></div>
  </div>
</div></body></html>""",

    "05_seleccion": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.back { color: #fff; font-size: 24px; margin-bottom: 16px; }
.page-title { font-size: 22px; font-weight: 700; color: #fff; margin-bottom: 20px; }
.activity-btn {
  background: #1E1E1E; border-radius: 16px; padding: 18px;
  display: flex; align-items: center; gap: 14px; margin-bottom: 10px;
}
.activity-icon {
  width: 44px; height: 44px; background: #00E6A820; border-radius: 12px;
  display: flex; align-items: center; justify-content: center; font-size: 22px;
}
.create-btn { color: #00E6A8; font-weight: 600; margin: 16px 0; font-size: 15px; }
.workout-item {
  background: #1E1E1E; border-radius: 12px; padding: 14px;
  display: flex; align-items: center; gap: 12px; margin-bottom: 8px;
}
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content">
    <div class="back">←</div>
    <div class="page-title">Iniciar una actividad</div>
    <div class="activity-btn"><div class="activity-icon">🏃</div><div class="white bold">Correr</div></div>
    <div class="activity-btn"><div class="activity-icon">🥾</div><div class="white bold">Senderismo</div></div>
    <div class="activity-btn"><div class="activity-icon">🚴</div><div class="white bold">Andar en bicicleta</div></div>
    <div class="activity-btn"><div class="activity-icon">🚶</div><div class="white bold">Caminar</div></div>
    <div class="section-title" style="margin-top:24px">Entrenamientos personalizados</div>
    <div class="create-btn">+ Crear entrenamiento</div>
    <div class="workout-item">
      <div class="activity-icon">🏃</div>
      <div><div class="white bold">Mi 5K matutino</div><div class="muted" style="font-size:13px">Correr · Meta: 5 km</div></div>
    </div>
    <div class="workout-item">
      <div class="activity-icon">🚴</div>
      <div><div class="white bold">Paseo en bici</div><div class="muted" style="font-size:13px">Bicicleta · Meta: 30 min</div></div>
    </div>
  </div>
</div></body></html>""",

    "06_tracking": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.tracking { text-align: center; padding-top: 40px; }
.activity-header { display: flex; align-items: center; justify-content: center; gap: 10px; margin-bottom: 8px; }
.duration { font-size: 64px; font-weight: 800; color: #fff; margin: 32px 0; font-variant-numeric: tabular-nums; }
.meta { color: #00E6A8; font-size: 14px; font-weight: 600; margin-bottom: 24px; }
.stats-row { display: flex; justify-content: space-around; margin: 32px 0; }
.stat-block .val { font-size: 24px; font-weight: 700; color: #fff; }
.stat-block .lbl { font-size: 12px; color: #888; margin-top: 4px; }
.controls { display: flex; justify-content: center; gap: 24px; margin-top: 40px; }
.ctrl-btn {
  width: 64px; height: 64px; border-radius: 50%; display: flex;
  align-items: center; justify-content: center; font-size: 24px;
}
.pause { background: #333; color: #fff; }
.stop { background: #FF5252; color: #fff; }
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content tracking">
    <div class="activity-header">
      <span style="font-size:24px">🏃</span>
      <span class="white bold" style="font-size:18px">Correr</span>
      <span style="margin-left:auto;font-size:20px">🔊</span>
    </div>
    <div class="meta">Meta: 5 km</div>
    <div class="duration">18:42</div>
    <div class="stats-row">
      <div class="stat-block"><div class="val">3.20</div><div class="lbl">Distancia (km)</div></div>
      <div class="stat-block"><div class="val">5:51</div><div class="lbl">Ritmo /km</div></div>
    </div>
    <div class="controls">
      <div class="ctrl-btn pause">⏸</div>
      <div class="ctrl-btn stop">⏹</div>
    </div>
  </div>
</div></body></html>""",

    "07_clima": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.weather-main { text-align: center; padding: 24px; }
.location { font-size: 22px; font-weight: 700; color: #fff; font-style: italic; }
.datetime { color: #888; font-size: 12px; margin-top: 4px; }
.temp-row { display: flex; align-items: center; justify-content: center; gap: 16px; margin: 16px 0; }
.temp { font-size: 64px; font-weight: 800; color: #fff; }
.condition { color: #ccc; font-size: 16px; }
.grid { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; margin-top: 16px; }
.grid-item { background: #1E1E1E; border-radius: 12px; padding: 14px; text-align: center; }
.grid-val { font-size: 18px; font-weight: 700; color: #fff; }
.grid-lbl { font-size: 11px; color: #888; margin-top: 4px; }
.recommendation {
  background: linear-gradient(135deg, #00E6A820, #1E1E1E);
  border: 1px solid #00E6A840; border-radius: 16px;
  padding: 16px; margin-top: 16px; color: #fff; font-size: 14px; line-height: 1.5;
}
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content">
    <div style="font-size:24px;color:#fff;margin-bottom:10px">←</div>
    <div class="card weather-main">
      <div class="location">Buenos Aires, AR</div>
      <div class="datetime">14:30 HS - 29 JUN</div>
      <div class="temp-row">
        <div class="temp">18°C</div>
        <div style="font-size:48px">⛅</div>
      </div>
      <div class="condition">Parcialmente nublado</div>
      <div class="grid">
        <div class="grid-item"><div class="grid-val">17°C</div><div class="grid-lbl">Sensación</div></div>
        <div class="grid-item"><div class="grid-val">65%</div><div class="grid-lbl">Humedad</div></div>
        <div class="grid-item"><div class="grid-val">12 KM/H</div><div class="grid-lbl">Viento</div></div>
        <div class="grid-item"><div class="grid-val">10 KM</div><div class="grid-lbl">Visibilidad</div></div>
      </div>
    </div>
    <div class="recommendation">
      🏃‍♂️ Clima fresco ideal para correr. Un rompevientos liviano es más que suficiente.
    </div>
  </div>
</div></body></html>""",

    "08_crear_entrenamiento": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.back { color: #fff; font-size: 24px; margin-bottom: 16px; }
.page-title { font-size: 22px; font-weight: 700; color: #fff; margin-bottom: 20px; }
.label { color: #fff; font-size: 14px; margin-top: 16px; margin-bottom: 6px; }
.radio-group { display: flex; flex-wrap: wrap; gap: 8px; margin-top: 8px; }
.radio {
  background: #1E1E1E; border: 1px solid #333; border-radius: 8px;
  padding: 10px 16px; color: #fff; font-size: 13px;
}
.radio.selected { border-color: #00E6A8; color: #00E6A8; }
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content">
    <div class="back">←</div>
    <div class="page-title">Crear entrenamiento</div>
    <div class="label">Nombre</div>
    <div class="input">Mi 5K matutino</div>
    <div class="label">Tipo de actividad</div>
    <div class="radio-group">
      <div class="radio selected">Correr</div>
      <div class="radio">Senderismo</div>
      <div class="radio">Bicicleta</div>
      <div class="radio">Caminar</div>
    </div>
    <div class="label">Meta (opcional)</div>
    <div class="radio-group">
      <div class="radio selected">Distancia (km)</div>
      <div class="radio">Tiempo (min)</div>
    </div>
    <div class="input" style="margin-top:12px">5</div>
    <div style="margin-top:32px" class="btn-primary">Guardar entrenamiento</div>
  </div>
</div></body></html>""",

    "09_mis_carreras": """<!DOCTYPE html><html><head><meta charset="utf-8"><style>
""" + COMMON_CSS + """
.back { color: #fff; font-size: 24px; }
.page-title { font-size: 22px; font-weight: 700; color: #fff; }
.header-row { display: flex; align-items: center; gap: 12px; margin-bottom: 20px; }
.run-card {
  background: #1E1E1E; border-radius: 12px; padding: 14px;
  display: flex; align-items: center; gap: 12px; margin-bottom: 8px;
}
.run-icon {
  width: 40px; height: 40px; background: #00E6A820; border-radius: 50%;
  display: flex; align-items: center; justify-content: center;
}
.btn-danger {
  border: 1px solid #FF5252; color: #FF5252; border-radius: 8px;
  padding: 12px; text-align: center; margin-top: 20px; font-size: 14px;
}
</style></head><body>
<div class="phone">
  <div class="status-bar"><span>9:41</span><span>📶 🔋</span></div>
  <div class="content">
    <div class="header-row">
      <div class="back">←</div>
      <div class="page-title">Mis carreras</div>
    </div>
    <div class="run-card">
      <div class="run-icon">🏃</div>
      <div style="flex:1">
        <div class="white bold">28/06/2026</div>
        <div class="muted" style="font-size:12px">5:45 /km · 32:10</div>
      </div>
      <div class="neon bold">5.60 km</div>
    </div>
    <div class="run-card">
      <div class="run-icon">🚴</div>
      <div style="flex:1">
        <div class="white bold">26/06/2026</div>
        <div class="muted" style="font-size:12px">3:12 /km · 48:00</div>
      </div>
      <div class="neon bold">15.00 km</div>
    </div>
    <div class="run-card">
      <div class="run-icon">🥾</div>
      <div style="flex:1">
        <div class="white bold">25/06/2026</div>
        <div class="muted" style="font-size:12px">8:20 /km · 1:05:00</div>
      </div>
      <div class="neon bold">7.80 km</div>
    </div>
    <div class="run-card">
      <div class="run-icon">🏃</div>
      <div style="flex:1">
        <div class="white bold">24/06/2026</div>
        <div class="muted" style="font-size:12px">6:10 /km · 25:30</div>
      </div>
      <div class="neon bold">4.15 km</div>
    </div>
    <div class="btn-danger">Borrar todas mis carreras</div>
  </div>
</div></body></html>""",
}


def write_mockups():
    MOCKUPS.mkdir(parents=True, exist_ok=True)
    for name, html in SCREENS.items():
        (MOCKUPS / f"{name}.html").write_text(html, encoding="utf-8")
    print(f"[OK] {len(SCREENS)} mockups escritos en {MOCKUPS}")


def take_screenshots():
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "playwright", "-q"])
        subprocess.check_call([sys.executable, "-m", "playwright", "install", "chromium"])
        from playwright.sync_api import sync_playwright

    SCREENSHOTS.mkdir(parents=True, exist_ok=True)
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(device_scale_factor=2)
        for name in SCREENS:
            html_path = MOCKUPS / f"{name}.html"
            page.goto(html_path.as_uri())
            page.set_viewport_size({"width": 430, "height": 900})
            out = SCREENSHOTS / f"{name}.png"
            page.locator(".phone").screenshot(path=str(out))
            print(f"  [IMG] {out.name}")
        browser.close()
    print(f"[OK] Capturas guardadas en {SCREENSHOTS}")


def build_word():
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml

    # Paleta corporativa Sprint
    C_PRIMARY = RGBColor(0, 230, 168)      # #00E6A8
    C_PRIMARY_DK = RGBColor(0, 160, 120)
    C_DARK = RGBColor(18, 18, 18)          # #121212
    C_DARK_MID = RGBColor(30, 30, 30)
    C_TEXT = RGBColor(45, 45, 45)
    C_MUTED = RGBColor(110, 110, 110)
    C_WHITE = RGBColor(255, 255, 255)
    C_ROW_ALT = "F4FAF8"
    C_HEADER_BG = "00C896"
    C_CALLOUT_BG = "E8FBF4"
    C_CALLOUT_BORDER = "00E6A8"

    doc = Document()
    figure_counter = [0]

    # ── Helpers XML / estilo ───────────────────────────────────────────────

    def set_cell_shading(cell, color_hex, foreground=False):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        if foreground:
            shading.set(qn("w:color"), "auto")
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tcPr.append(mar)

    def remove_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tblPr.append(borders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def set_table_borders(table, color="D0D0D0", size="4"):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), color)
            borders.append(el)
        tblPr.append(borders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def style_run(run, *, size=11, bold=False, italic=False, color=C_TEXT, font="Calibri"):
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rPr.insert(0, rFonts)

    def add_horizontal_rule(space_before=6, space_after=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "00E6A8")
        bottom.set(qn("w:space"), "1")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_page_number_field(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        style_run(run, size=9, color=C_MUTED)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        run2 = paragraph.add_run()
        style_run(run2, size=9, color=C_MUTED)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)
        run3 = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_sep)
        run4 = paragraph.add_run("1")
        style_run(run4, size=9, color=C_MUTED)
        run5 = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fld_end)

    def setup_page(section, *, header=True):
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        if not header:
            return
        header_el = section.header
        header_el.is_linked_to_previous = False
        hp = header_el.paragraphs[0]
        hp.clear()
        run_left = hp.add_run("SPRINT  |  Informe Tecnico")
        style_run(run_left, size=8, color=C_MUTED, bold=True)
        hp.paragraph_format.space_after = Pt(0)
        footer_el = section.footer
        footer_el.is_linked_to_previous = False
        fp = footer_el.paragraphs[0]
        fp.clear()
        run_f = fp.add_run("RunTracker — Android  |  Junio 2026    ")
        style_run(run_f, size=8, color=C_MUTED)
        add_page_number_field(fp)

    def add_section_title(text, level=1):
        if level == 1:
            add_horizontal_rule(12, 4)
        h = doc.add_heading(level=level)
        h.clear()
        run = h.add_run(text)
        if level == 1:
            style_run(run, size=18, bold=True, color=C_PRIMARY_DK, font="Calibri Light")
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(10)
        else:
            style_run(run, size=13, bold=True, color=C_DARK, font="Calibri")
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        return h

    def add_body(text, *, size=11, color=C_TEXT, space_after=8, justify=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        style_run(run, size=size, color=color)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.25
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(text)
        style_run(run, size=10.5, color=C_TEXT)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.2
        return p

    def add_callout(title, body):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, C_CALLOUT_BG)
        set_cell_margins(cell, 140, 140, 180, 180)
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "12")
            el.set(qn("w:color"), C_CALLOUT_BORDER)
            borders.append(el)
        tcPr.append(borders)
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(title)
        style_run(r1, size=11, bold=True, color=C_PRIMARY_DK)
        p1.paragraph_format.space_after = Pt(6)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(body)
        style_run(r2, size=10.5, color=C_TEXT)
        p2.paragraph_format.line_spacing = 1.25
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_toc_entry(text, page_hint=""):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, leader=1)
        run = p.add_run(text)
        style_run(run, size=11, color=C_TEXT)
        if page_hint:
            p.add_run("\t")
            run2 = p.add_run(page_hint)
            style_run(run2, size=11, color=C_MUTED)
        p.paragraph_format.space_after = Pt(5)

    def add_styled_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="C8E8DC")
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            r = p.add_run(h)
            style_run(r, size=10, bold=True, color=C_WHITE)
            set_cell_shading(hdr_cells[i], C_HEADER_BG)
            set_cell_margins(hdr_cells[i])
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for idx, row_data in enumerate(rows):
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = ""
                p = row[i].paragraphs[0]
                r = p.add_run(str(val))
                bold = i == 0
                style_run(r, size=9.5, bold=bold, color=C_TEXT)
                if idx % 2 == 1:
                    set_cell_shading(row[i], C_ROW_ALT)
                set_cell_margins(row[i], 60, 60, 100, 100)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        return table

    def add_screen_section(img_file, section_title, description, activity_name):
        figure_counter[0] += 1
        fig_num = figure_counter[0]
        add_section_title(section_title, level=2)

        table = doc.add_table(rows=1, cols=2)
        remove_table_borders(table)
        left, right = table.rows[0].cells[0], table.rows[0].cells[1]
        set_cell_margins(left, 0, 0, 0, 160)
        set_cell_margins(right, 0, 0, 0, 0)
        left.width = Inches(3.55)
        right.width = Inches(2.35)

        p_desc = left.paragraphs[0]
        r_desc = p_desc.add_run(description)
        style_run(r_desc, size=10.5, color=C_TEXT)
        p_desc.paragraph_format.line_spacing = 1.3
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p_tag = left.add_paragraph()
        p_tag.paragraph_format.space_before = Pt(10)
        r_tag = p_tag.add_run(f"Activity: {activity_name}")
        style_run(r_tag, size=9, bold=True, color=C_PRIMARY_DK)

        img_path = SCREENSHOTS / img_file
        if img_path.exists():
            p_img = right.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(str(img_path), width=Inches(2.05))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(16)
        r_cap = cap.add_run(f"Figura {fig_num} — {activity_name}")
        style_run(r_cap, size=9, italic=True, color=C_MUTED)

    def add_kpi_row(items):
        """items: list of (value, label)"""
        table = doc.add_table(rows=1, cols=len(items))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)
        for i, (val, lbl) in enumerate(items):
            cell = table.rows[0].cells[i]
            set_cell_shading(cell, "1A1D21")
            set_cell_margins(cell, 120, 120, 80, 80)
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(val)
            style_run(r1, size=16, bold=True, color=C_PRIMARY)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(lbl)
            style_run(r2, size=8.5, color=C_MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── Configurar secciones ─────────────────────────────────────────────────
    setup_page(doc.sections[0], header=False)

    # ══════════════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════════════
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover.rows[0].cells[0]
    set_cell_shading(cover_cell, "121212")
    set_cell_margins(cover_cell, 600, 600, 400, 400)
    cover_cell.height = Cm(24)

    # Logo
    p_logo = cover_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(20)
    r_logo = p_logo.add_run("  S  ")
    style_run(r_logo, size=28, bold=True, color=RGBColor(0, 0, 0), font="Calibri")
    # Simular circulo con fondo en run shading no es trivial; usamos texto estilizado
    r_logo.font.highlight_color = None

    p_brand = cover_cell.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_brand = p_brand.add_run("SPRINT")
    style_run(r_brand, size=44, bold=True, color=C_PRIMARY, font="Calibri Light")
    p_brand.paragraph_format.space_after = Pt(8)

    p_tag = cover_cell.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tag = p_tag.add_run("Tu companero de carrera")
    style_run(r_tag, size=12, color=RGBColor(160, 160, 160))
    p_tag.paragraph_format.space_after = Pt(36)

    p_doc_title = cover_cell.add_paragraph()
    p_doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dt = p_doc_title.add_run("INFORME TECNICO")
    style_run(r_dt, size=14, bold=True, color=C_WHITE)
    p_doc_title.paragraph_format.space_after = Pt(4)

    p_doc_sub = cover_cell.add_paragraph()
    p_doc_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ds = p_doc_sub.add_run("Documentacion de la Aplicacion Movil RunTracker")
    style_run(r_ds, size=11, color=RGBColor(180, 180, 180))
    p_doc_sub.paragraph_format.space_after = Pt(40)

    # Metadata portada
    meta = cover_cell.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta)
    meta_data = [
        ("Plataforma", "Android (Java 11)"),
        ("Version", "1.0 — com.example.sprint"),
        ("SDK", "Min 24 / Target 36"),
        ("Fecha", "Junio 2026"),
    ]
    for i, (k, v) in enumerate(meta_data):
        c0, c1 = meta.rows[i].cells[0], meta.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        style_run(r0, size=9, bold=True, color=C_PRIMARY)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = c1.paragraphs[0].add_run(v)
        style_run(r1, size=9, color=RGBColor(200, 200, 200))
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_margins(c0, 40, 40, 80, 20)
        set_cell_margins(c1, 40, 40, 20, 80)

    p_line = cover_cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(30)
    r_line = p_line.add_run("_" * 40)
    style_run(r_line, size=8, color=C_PRIMARY)

    doc.add_page_break()

    # Nueva seccion con header/footer
    new_sec = doc.add_section()
    setup_page(new_sec, header=True)

    # ══════════════════════════════════════════════════════════════════════════
    # INDICE
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("Indice de contenidos", 1)
    add_body(
        "Este documento describe de forma integral la aplicacion Sprint (RunTracker): "
        "arquitectura, tecnologias, pantallas, modelo de datos y flujos de usuario.",
        space_after=14,
    )
    toc_items = [
        "1. Resumen ejecutivo",
        "2. Introduccion y objetivos",
        "3. Descripcion general",
        "4. Stack tecnologico",
        "5. Arquitectura del sistema",
        "6. Pantallas y funcionalidades",
        "7. Modelo de datos (Firebase Firestore)",
        "8. Flujo de navegacion",
        "9. Componentes principales",
        "10. Integraciones externas",
        "11. Permisos y seguridad",
        "12. Conclusiones y trabajo futuro",
    ]
    for item in toc_items:
        add_toc_entry(item)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("1. Resumen ejecutivo", 1)
    add_kpi_row([
        ("10", "Pantallas"),
        ("4", "Actividades"),
        ("3", "Capas"),
        ("2", "APIs externas"),
    ])
    add_callout(
        "Vision del proyecto",
        "Sprint es una aplicacion Android nativa orientada al fitness que combina registro de "
        "entrenamientos, estadisticas motivacionales, perfil de usuario con datos de salud y "
        "pronostico meteorologico. Los datos se sincronizan en Firebase Firestore y la "
        "autenticacion se gestiona con Firebase Auth.",
    )
    add_body(
        "El producto responde a la necesidad de centralizar la actividad fisica del usuario en "
        "una interfaz moderna de tema oscuro, con metricas diarias (pasos, distancia, calorias, "
        "tiempo y racha), entrenamientos personalizables y un historial persistente en la nube.",
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. INTRODUCCION
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("2. Introduccion y objetivos", 1)
    add_body(
        "Sprint (nombre comercial de la app RunTracker) es una aplicacion movil nativa para Android "
        "desarrollada en Java, orientada al seguimiento de actividades fisicas como correr, caminar, "
        "senderismo y ciclismo. Permite registrar entrenamientos, visualizar estadisticas, gestionar "
        "el perfil personal y consultar el clima antes de salir a entrenar.",
    )
    add_body("Objetivos principales del proyecto:")
    objectives = [
        "Dashboard motivacional con metricas del dia: pasos, distancia, calorias, tiempo y racha.",
        "Registro y seguimiento de actividades fisicas en tiempo real con pausa y reanudacion.",
        "Sincronizacion de datos del usuario y carreras en la nube mediante Firebase.",
        "Entrenamientos personalizados con metas de distancia (km) o tiempo (minutos).",
        "Integracion meteorologica con recomendaciones segun temperatura ambiente.",
    ]
    for obj in objectives:
        add_bullet(obj)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DESCRIPCION GENERAL
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("3. Descripcion general de la aplicacion", 1)
    add_body(
        "La aplicacion adopta un diseno oscuro (#121212) con acentos en verde neon (#00E6A8), "
        "inspirado en apps de fitness modernas. La navegacion principal usa una barra inferior "
        "con acceso a Inicio, Perfil y un boton flotante central (FAB) para iniciar actividades.",
    )
    add_styled_table(
        ["Atributo", "Valor"],
        [
            ("Nombre comercial", "Sprint"),
            ("Package ID", "com.example.sprint"),
            ("Version", "1.0 (versionCode 1)"),
            ("Lenguaje", "Java 11"),
            ("SDK minimo", "24 — Android 7.0 Nougat"),
            ("SDK objetivo", "36"),
            ("Tema UI", "Material Design 3 — modo oscuro"),
            ("Persistencia", "Firebase Firestore + SharedPreferences"),
        ],
        col_widths=[2.2, 4.3],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 4. STACK
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("4. Stack tecnologico", 1)
    add_styled_table(
        ["Tecnologia", "Version / Detalle", "Rol en el proyecto"],
        [
            ("Java", "11", "Lenguaje principal"),
            ("Android SDK", "36", "Plataforma movil"),
            ("Material Design", "1.13.0", "Componentes UI"),
            ("Firebase Auth", "BOM 33.8", "Autenticacion email/contrasena"),
            ("Cloud Firestore", "BOM 33.8", "Base de datos NoSQL"),
            ("Retrofit", "2.9.0", "Cliente HTTP REST"),
            ("Gson", "2.9.0", "Deserializacion JSON"),
            ("Play Services Location", "21.3.0", "Servicios de ubicacion"),
            ("MPAndroidChart", "v3.1.0", "Graficos (dependencia)"),
            ("SharedPreferences", "Nativo", "Cache local de perfil"),
        ],
        col_widths=[1.8, 1.4, 3.3],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 5. ARQUITECTURA
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("5. Arquitectura del sistema", 1)
    add_body(
        "La aplicacion implementa una arquitectura por capas que separa presentacion, logica de "
        "negocio y acceso a datos. El patron Repository desacopla las Activities de Firestore.",
    )
    add_styled_table(
        ["Capa", "Componentes", "Responsabilidad"],
        [
            ("Presentacion", "10 Activities + layouts XML", "UI, eventos, navegacion"),
            ("Logica", "Repositories, Stopwatch, Stats", "Reglas de negocio y calculos"),
            ("Datos", "Firestore, SharedPreferences, files/", "Persistencia local y remota"),
        ],
        col_widths=[1.5, 2.5, 2.5],
    )
    add_callout(
        "Patron de diseno",
        "RunRepository y WorkoutRepository centralizan el acceso a Firestore. Los modelos Run y "
        "Workout son inmutables y se construyen desde DocumentSnapshot mediante factory methods "
        "(fromDocument), garantizando un contrato de datos consistente en toda la app.",
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. PANTALLAS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("6. Pantallas y funcionalidades", 1)
    add_body(
        "A continuacion se documentan las diez pantallas de la aplicacion. Cada una incluye "
        "descripcion funcional, clase Java asociada y captura de la interfaz.",
        space_after=12,
    )

    screens_info = [
        ("01_splash.png", "6.1 Splash", "SplashActivity",
         "Pantalla de arranque con duracion de 1,5 segundos. Consulta Firebase Auth "
         "y redirige a MainActivity si hay sesion activa, o a LoginActivity en caso contrario."),
        ("02_login.png", "6.2 Autenticacion", "LoginActivity",
         "Login y registro por email/contrasena. Incluye recuperacion de contrasena, botones "
         "sociales (Google/GitHub planificados) y creacion del documento de usuario en Firestore al registrarse."),
        ("03_dashboard.png", "6.3 Dashboard", "MainActivity",
         "Pantalla principal: saludo personalizado, frase motivacional, contador de pasos, "
         "objetivo diario (10.000 pasos), racha, acceso al clima y metricas del dia."),
        ("04_perfil.png", "6.4 Perfil", "ProfileActivity",
         "Avatar, datos de salud (peso, altura, IMC), estadisticas totales, logros "
         "desbloqueables, preview de carreras, edicion de perfil y cierre de sesion."),
        ("05_seleccion.png", "6.5 Seleccion de actividad", "WorkoutSelectionActivity",
         "Cuatro tipos de actividad rapidos y lista de entrenamientos personalizados "
         "guardados en Firestore. Acceso al formulario de creacion."),
        ("06_tracking.png", "6.6 Tracking en vivo", "WorkoutTrackingActivity",
         "Cronometro con pausa/reanudacion, distancia simulada segun tipo de actividad, "
         "ritmo por km y guardado automatico en Firestore al detener."),
        ("07_clima.png", "6.7 Clima", "WeatherActivity",
         "Datos de OpenWeatherMap para Buenos Aires: temperatura, humedad, viento, "
         "visibilidad y recomendacion de entrenamiento segun condiciones."),
        ("08_crear_entrenamiento.png", "6.8 Crear entrenamiento", "CreateWorkoutActivity",
         "Formulario con nombre, tipo de actividad y meta opcional (distancia o tiempo). "
         "Persiste en la subcoleccion workouts del usuario."),
        ("09_mis_carreras.png", "6.9 Historial", "MyRunsActivity",
         "Listado completo de carreras con fecha, ritmo, duracion y distancia. "
         "Permite borrar todo el historial con dialogo de confirmacion."),
    ]

    for img, title, activity, desc in screens_info:
        add_screen_section(img, title, desc, activity)

    doc.add_page_break()

    # Galeria compacta 2x2 de pantallas principales
    add_section_title("6.10 Galeria de pantallas principales", 2)
    gallery = [
        ("03_dashboard.png", "Dashboard"),
        ("04_perfil.png", "Perfil"),
        ("06_tracking.png", "Tracking"),
        ("07_clima.png", "Clima"),
    ]
    gtable = doc.add_table(rows=2, cols=2)
    remove_table_borders(gtable)
    for i, (img, label) in enumerate(gallery):
        row, col = divmod(i, 2)
        cell = gtable.rows[row].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = SCREENSHOTS / img
        if img_path.exists():
            p.add_run().add_picture(str(img_path), width=Inches(1.85))
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(label)
        style_run(r, size=9, bold=True, color=C_MUTED)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. MODELO DE DATOS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("7. Modelo de datos (Firebase Firestore)", 1)
    add_body("Estructura jerarquica de colecciones y campos:")
    add_styled_table(
        ["Ruta", "Campo", "Tipo", "Descripcion"],
        [
            ("users/{uid}", "username", "string", "Nombre del usuario"),
            ("users/{uid}", "joinYear", "number", "Ano de registro"),
            ("users/{uid}", "weightKg", "number", "Peso en kilogramos"),
            ("users/{uid}", "heightCm", "number", "Altura en centimetros"),
            ("users/{uid}", "age", "number", "Edad"),
            ("users/{uid}", "gender", "string", "Codigo M / F / O"),
            (".../runs/{id}", "tipo", "string", "run | hike | bike | walk"),
            (".../runs/{id}", "distanciaKm", "number", "Distancia recorrida"),
            (".../runs/{id}", "pasos", "number", "Cantidad de pasos"),
            (".../runs/{id}", "duracionSegundos", "number", "Duracion total"),
            (".../runs/{id}", "timestamp", "timestamp", "Fecha del servidor"),
            (".../workouts/{id}", "name", "string", "Nombre del entrenamiento"),
            (".../workouts/{id}", "type", "string", "Tipo de actividad"),
            (".../workouts/{id}", "goalType", "string", "distance | time | vacio"),
            (".../workouts/{id}", "goalValue", "number", "Km o minutos"),
        ],
        col_widths=[1.6, 1.3, 0.9, 2.7],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 8. FLUJO
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("8. Flujo de navegacion", 1)
    add_styled_table(
        ["Origen", "Accion", "Destino"],
        [
            ("SplashActivity", "Sin sesion", "LoginActivity"),
            ("SplashActivity", "Con sesion", "MainActivity"),
            ("LoginActivity", "Auth exitoso", "MainActivity"),
            ("MainActivity", "Navbar Perfil", "ProfileActivity"),
            ("MainActivity", "FAB / Correr", "WorkoutSelectionActivity"),
            ("MainActivity", "Card clima", "WeatherActivity"),
            ("ProfileActivity", "Editar", "EditProfileActivity"),
            ("ProfileActivity", "Ver carreras", "MyRunsActivity"),
            ("WorkoutSelection", "Iniciar actividad", "WorkoutTrackingActivity"),
            ("WorkoutSelection", "Crear", "CreateWorkoutActivity"),
        ],
        col_widths=[2.0, 1.8, 2.7],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 9. COMPONENTES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("9. Componentes y clases principales", 1)
    add_styled_table(
        ["Clase", "Tipo", "Funcion"],
        [
            ("ActivityType", "enum", "Catalogo de actividades con velocidad simulada"),
            ("Stopwatch", "util", "Cronometro con pausa real (SystemClock)"),
            ("RunRepository", "repository", "CRUD de carreras en Firestore"),
            ("WorkoutRepository", "repository", "CRUD de entrenamientos"),
            ("UserPreferences", "util", "Wrapper de SharedPreferences"),
            ("ImageStorage", "util", "Foto de perfil local (512px, JPEG 85%)"),
            ("DashboardStats", "model", "Metricas del dia y racha"),
            ("RunStats", "model", "Agregados: km, ritmo, max distancia"),
            ("RunCardFactory", "factory", "Tarjetas de carrera reutilizables"),
            ("NavbarHelper", "util", "Estado activo de tabs"),
            ("RetrofitClient", "api", "Cliente singleton OpenWeatherMap"),
        ],
        col_widths=[1.8, 1.0, 3.7],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 10. INTEGRACIONES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("10. Integraciones externas", 1)

    add_section_title("10.1 Firebase (Google)", 2)
    add_bullet("Firebase Authentication: registro, login, logout y reset de contrasena.")
    add_bullet("Cloud Firestore: perfiles, carreras y entrenamientos por usuario autenticado.")

    add_section_title("10.2 OpenWeatherMap", 2)
    add_styled_table(
        ["Parametro", "Valor"],
        [
            ("Endpoint", "GET /data/2.5/weather"),
            ("Base URL", "api.openweathermap.org"),
            ("Unidades", "metric (Celsius)"),
            ("Idioma", "es (espanol)"),
            ("Ubicacion default", "Buenos Aires (-34.60, -58.38)"),
        ],
        col_widths=[2.0, 4.5],
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 11. PERMISOS
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("11. Permisos y seguridad", 1)
    add_styled_table(
        ["Permiso", "Proposito"],
        [
            ("ACCESS_FINE_LOCATION", "Ubicacion precisa para tracking"),
            ("ACCESS_COARSE_LOCATION", "Ubicacion aproximada"),
            ("INTERNET", "Firebase y API meteorologica"),
        ],
        col_widths=[2.5, 4.0],
    )
    add_callout(
        "Consideraciones de seguridad",
        "La autenticacion se delega a Firebase Auth. Las reglas de Firestore deben limitar "
        "lectura/escritura al propio uid del usuario. La foto de perfil se almacena solo en "
        "el almacenamiento interno de la app, sin subir a la nube.",
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 12. CONCLUSIONES
    # ══════════════════════════════════════════════════════════════════════════
    add_section_title("12. Conclusiones y trabajo futuro", 1)
    add_body(
        "Sprint demuestra la integracion de tecnologias moviles modernas: autenticacion en la "
        "nube, base de datos NoSQL, consumo de APIs REST, Material Design y persistencia hibrida "
        "(local + remota). La arquitectura por capas con repositorios facilita el mantenimiento "
        "y la evolucion del producto.",
    )
    add_body("Lineas de mejora identificadas:")
    improvements = [
        "Tracking GPS real con Google Play Services Location.",
        "Login social con Google y GitHub.",
        "Subida de foto de perfil a Firebase Storage.",
        "Notificaciones push para recordatorios de entrenamiento.",
        "Graficos de progreso con MPAndroidChart.",
        "Modo offline con sincronizacion automatica.",
    ]
    for imp in improvements:
        add_bullet(imp)

    # Cierre
    doc.add_paragraph()
    add_horizontal_rule(20, 8)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("— Fin del informe —")
    style_run(r_end, size=10, italic=True, color=C_MUTED)

    doc.save(str(OUTPUT))
    print(f"\n[OK] Informe generado: {OUTPUT}")


if __name__ == "__main__":
    import sys
    word_only = "--word-only" in sys.argv
    screenshots_ready = all((SCREENSHOTS / f"{name}.png").exists() for name in SCREENS)

    if word_only or screenshots_ready:
        if not screenshots_ready:
            print("Generando mockups y capturas primero...")
            write_mockups()
            take_screenshots()
        else:
            print("Capturas existentes encontradas, regenerando solo el Word...")
        print("\nConstruyendo documento Word...")
        build_word()
    else:
        print("=== Generando informe Sprint ===\n")
        write_mockups()
        print("\nTomando capturas...")
        take_screenshots()
        print("\nConstruyendo documento Word...")
        build_word()
    print("\nListo!")
